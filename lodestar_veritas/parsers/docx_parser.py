from pathlib import Path

from docx import Document


class DOCXParser:
    """
    Real DOCX parser.

    Extracts:
    - paragraphs
    - headings
    - tables
    """

    def parse(self, file_path: str) -> dict:
        path = Path(file_path)
        document = Document(file_path)

        text_parts = []
        paragraphs = []
        tables = []

        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name if paragraph.style else ""

            paragraph_data = {
                "paragraph_number": index,
                "text": text,
                "style": style_name,
            }

            paragraphs.append(paragraph_data)
            text_parts.append(f"[Paragraph {index} | {style_name}]\n{text}")

        for table_index, table in enumerate(document.tables, start=1):
            table_rows = []

            for row in table.rows:
                row_values = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_values)

            tables.append(
                {
                    "table_number": table_index,
                    "rows": table_rows,
                }
            )

            text_parts.append(f"[Table {table_index}]")
            for row in table_rows:
                text_parts.append(" | ".join(row))

        content = "\n\n".join(text_parts).strip()

        return {
            "content": content,
            "paragraphs": paragraphs,
            "tables": tables,
            "metadata": {
                "source": file_path,
                "file_name": path.name,
                "parser": "DOCXParser",
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
            },
        }