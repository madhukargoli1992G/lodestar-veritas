from pathlib import Path

from docx import Document

from lodestar_veritas.parsers.docx_parser import DOCXParser


def test_docx_parser_extracts_paragraphs_and_tables(tmp_path: Path):
    docx_path = tmp_path / "sample.docx"

    document = Document()
    document.add_heading("Risk Factors", level=1)
    document.add_paragraph("Revenue increased by 25% in 2024.")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "25%"

    document.save(docx_path)

    parser = DOCXParser()
    result = parser.parse(str(docx_path))

    assert result["metadata"]["parser"] == "DOCXParser"
    assert result["metadata"]["paragraph_count"] >= 2
    assert result["metadata"]["table_count"] == 1
    assert "Revenue increased" in result["content"]